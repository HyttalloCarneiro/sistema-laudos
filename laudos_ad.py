import openai  # Certifique-se de ter instalado com: pip install openai
from docx import Document  # Certifique-se de ter instalado o python-docx
from docx.shared import Inches
import base64
from io import BytesIO
import streamlit as st
from datetime import datetime, date
from PyPDF2 import PdfReader
import os
import re

# Lista global de CIDs para referência de patologias
lista_cids = {
    "M50": "Transtorno de disco intervertebral cervical",
    "M51": "Transtorno de disco intervertebral lombar",
    "M79.7": "Fibromialgia",
    # Adicione outros CIDs conforme necessário
}

def extrair_dados_pdf(pdf_file):
    try:
        reader = PdfReader(pdf_file)
        texto = ""
        for page in reader.pages[:3]:
            texto += page.extract_text()

        # --- OCR nas páginas 14 e 15 (índices 13 e 14) ---
        import pytesseract
        from PIL import Image
        import fitz  # PyMuPDF
        try:
            # Reposiciona ponteiro do arquivo para o início antes de fitz.open
            pdf_file.seek(0)
            doc_fitz = fitz.open(stream=pdf_file.read(), filetype="pdf")
            for i in [13, 14]:
                if i < len(doc_fitz):
                    page = doc_fitz.load_page(i)
                    pix = page.get_pixmap(dpi=300)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    texto_ocr = pytesseract.image_to_string(img, lang='por')
                    texto += "\n" + texto_ocr
            pdf_file.seek(0)  # Reposiciona ponteiro do arquivo
        except Exception as ocr_err:
            texto += f"\n[Erro OCR páginas 14-15: {ocr_err}]"

        return texto
    except Exception as e:
        return f"Erro ao processar PDF: {e}"

# Extrai nome, CPF, RG, DER, NB e data de nascimento a partir do texto extraído do PDF
def extrair_dados_estruturados(texto):
    nome = re.search(r"Terceiro Vinculado\s+([^\(]+)", texto)
    cpf = re.search(r"CPF(?: Nº)?:\s*([\d\.\-]+)", texto, re.IGNORECASE)
    rg = re.search(r"RG(?: Nº)?:\s*([\w\/\.\-]+)", texto, re.IGNORECASE)
    profissao = re.search(r"Profiss[aã]o:\s*(.+)", texto, re.IGNORECASE)
    der = re.search(r"DER:\s*(\d{2}/\d{2}/\d{4})", texto, re.IGNORECASE)
    if not der:
        der = re.search(r"Data de Entrada do Requerimento\s*[:\-]?\s*(\d{2}/\d{2}/\d{4})", texto, re.IGNORECASE)
    nascimento = re.search(r"Nasc(?:imento)?:\s*(\d{2}/\d{2}/\d{4})", texto, re.IGNORECASE)
    nb = re.search(r"NB(?:\.|:)?\s*(\d{3}\.\d{5}\.\d{2}|\d{10})", texto, re.IGNORECASE)

    # Buscas alternativas para RG, CPF e data de nascimento se não encontrados pelos padrões acima
    if not nascimento:
        nascimento = re.search(r"(\d{2}/\d{2}/\d{4})", texto)
    if not rg:
        rg = re.search(r"RG[^:\d]{0,6}[:\s]*([\w\/\.\-]{5,})", texto)
    if not cpf:
        cpf = re.search(r"CPF[^:\d]{0,6}[:\s]*([\d\.\-]{11,})", texto)

    beneficio_anterior = ""
    match = re.search(r"(concedido|indeferido|nega[du]?)[\s\w,]*benef[ií]cio", texto, re.IGNORECASE)
    if match:
        beneficio_anterior = match.group(0).capitalize()

    return {
        "nome": nome.group(1).strip() if nome else "",
        "cpf": cpf.group(1).strip() if cpf else "",
        "rg": rg.group(1).strip() if rg else "",
        "profissao": profissao.group(1).strip() if profissao else "",
        "der": der.group(1).strip() if der else "",
        "nascimento": nascimento.group(1).strip() if nascimento else "",
        "nb": nb.group(1).strip() if nb else "",
        "beneficio_anterior": beneficio_anterior,
    }

def calcular_idade(data_nasc_str):
    try:
        dia, mes, ano = map(int, data_nasc_str.split('/'))
        nascimento = date(ano, mes, dia)
        hoje = date.today()
        idade = hoje.year - nascimento.year - ((hoje.month, hoje.day) < (nascimento.month, nascimento.day))
        return idade
    except:
        return None

def interface_laudo_ad():
    # Inicializar valores no session_state se ainda não existem
    if "numero_processo_manual" not in st.session_state:
        st.session_state["numero_processo_manual"] = ""
    if "nome_autor_manual" not in st.session_state:
        st.session_state["nome_autor_manual"] = ""
    st.markdown("### 📂 Carregar Processo (PDF)")
    col_num_proc, col_nome_autor = st.columns(2)
    with col_num_proc:
        numero_input = st.text_input(
            "Número do processo:",
            value=st.session_state.get("numero_processo_manual", ""),
            key="numero_processo_manual"
        )
        # st.session_state["numero_processo_manual"] = numero_input  # Removido conforme instrução
        # Validação de formato
        if numero_input and not re.match(r"^\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}$", numero_input):
            st.warning("⚠️ O número do processo inserido não está no formato esperado: 0000000-00.0000.0.00.0000")
    with col_nome_autor:
        nome_input = st.text_input(
            "Nome do(a) autor(a):",
            value=st.session_state.get("nome_autor_manual", ""),
            key="nome_autor_manual"
        )
        # st.session_state["nome_autor_manual"] = nome_input  # Removido conforme instrução
    processo_pdf = st.file_uploader("Selecione o arquivo do processo completo em PDF", type=["pdf"], key="processo_pdf")
    if processo_pdf is not None and "texto_extraido" not in st.session_state:
        texto_extraido = extrair_dados_pdf(processo_pdf)
        st.session_state["texto_extraido"] = texto_extraido
        st.session_state["processo_carregado_com_sucesso"] = True
        st.rerun()
    # Nova verificação para exibir imediatamente a tela de laudo
    if st.session_state.get("processo_carregado_com_sucesso"):
        st.session_state["processo_carregado_com_sucesso"] = False
        st.rerun()
    # Sincronizar lista_cids com patologias_memoria, se existir
    if "patologias_memoria" in st.session_state:
        for cid, dados_pat in st.session_state["patologias_memoria"].items():
            lista_cids[cid] = dados_pat["descricao"]
            if "conceitos" not in st.session_state:
                st.session_state["conceitos"] = {}
            st.session_state["conceitos"][cid] = dados_pat["conceito"]
    # Novo bloco para carregar texto extraído do processo
    texto_extraido = st.session_state.get("texto_extraido", "")
    if not texto_extraido:
        st.warning("Nenhum processo foi carregado. Volte à tela anterior para selecionar um processo.")
        return

    dados = extrair_dados_estruturados(texto_extraido)
    st.session_state["dados_extraidos"] = dados

    numero_processo = st.session_state.get("numero_processo_manual", "Não encontrado")

    nome_autor = st.session_state.get("nome_autor_manual", "") or st.session_state.get("dados_extraidos", {}).get("nome", "")
    data_nascimento = st.session_state.get("dados_extraidos", {}).get("nascimento", "")
    idade = calcular_idade(data_nascimento) if data_nascimento else None
    historico_beneficios = dados.get("beneficio_anterior", "Nega benefício previdenciário prévio.")

    # Barra lateral - ordem e informações conforme solicitado
    st.sidebar.button("⬅️ Voltar")
    if st.session_state.get("tem_quesitos_autor", False):
        st.sidebar.markdown('<p style="color:red; font-weight:bold;">QUESITOS</p>', unsafe_allow_html=True)
    st.sidebar.markdown("### Dados do Periciando Extraídos")
    st.sidebar.markdown(f"**Número do processo:** {numero_processo}")
    st.sidebar.markdown(f"**Nome completo:** {nome_autor if nome_autor else 'Não informado'}")
    st.sidebar.markdown(f"**Data de nascimento:** {data_nascimento if data_nascimento else 'Não informado'}")
    st.sidebar.markdown(f"**Idade:** {idade if idade is not None else 'Não informado'} anos")
    st.sidebar.markdown(f"**Tipo de benefício:** Auxílio-Doença")
    der_extraida = st.session_state.get('dados_extraidos', {}).get('der', '')
    st.session_state["der_editavel"] = st.text_input("DER:", value=der_extraida, key="input_der")
    st.sidebar.markdown(f"**DER:** {st.session_state['der_editavel']}")
    st.session_state["historico_beneficios"] = st.text_input("Situação benefício anterior:", value=historico_beneficios, key="input_hist_benef")
    st.sidebar.markdown(f"**Situação benefício anterior:** {st.session_state['historico_beneficios']}")

    # Imagens: mover uploaders para a barra lateral
    st.sidebar.markdown("### 📷 Imagens")
    col_foto, col_docs = st.sidebar.columns(2)
    with col_foto:
        foto_autor = st.file_uploader("Foto 3x4", type=["jpg", "jpeg", "png"], key="foto_autor")
    with col_docs:
        fotos_anexos = st.file_uploader("Anexos", type=["jpg", "jpeg", "png"], accept_multiple_files=True, key="anexos_docs")

    # Removido subheader e uploader comentado conforme instruções

    st.markdown(f"## 👤 {nome_autor}")

    st.markdown('<p style="font-size:16px; font-weight:bold;">Profissão:</p>', unsafe_allow_html=True)
    profissao = st.text_input("Profissão declarada", value=st.session_state.get("dados_extraidos", {}).get("profissao", "Rurícola"))

    st.markdown('<p style="font-size:16px; font-weight:bold;">Histórico Laboral:</p>', unsafe_allow_html=True)
    historico_laboral = st.text_area("Outras ocupações já exercidas pelo periciando:")

    st.markdown('<p style="font-size:16px; font-weight:bold;">Escolaridade:</p>', unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1:
        esc_analfabeto = st.checkbox("Analfabeto")
        esc_fund_inc = st.checkbox("Fundamental incompleto")
        esc_fund_comp = st.checkbox("Fundamental completo")
    with col2:
        esc_so_assina = st.checkbox("Só assina o nome")
        esc_med_inc = st.checkbox("Médio incompleto")
        esc_med_comp = st.checkbox("Médio completo")
    with col3:
        esc_nao_sabe = st.checkbox("Não sabe informar")
        esc_sup_inc = st.checkbox("Superior incompleto")
        esc_sup_comp = st.checkbox("Superior completo")

    # Substituir bloco "🧠 3. Anamnese e Exame Médico"
    st.subheader("🧠 Anamnese")
    anamnese_default = "Durante ato pericial a parte autora relata que iniciou quadro de crise convulsões há mais de 20 anos (DID em 2004) tendo iniciado uso de medicações. Refere que evolui com crises frequentes, com episódios semanais, vindo também com crises de esquecimento frequentes, não vindo conseguindo realizar suas atividades. Atualmente em uso de Carbamazepina 200mg 2x ao die e Levetiracetam 500mg 2x ao dia."
    anamnese = st.text_area("Descreva a Anamnese:", value=anamnese_default if 'anamnese' not in st.session_state else st.session_state.anamnese, height=200)
    st.session_state.anamnese = anamnese

    st.subheader("🩺 Exame Clínico")
    exame_clinico_default = "Ao exame clínico se apresenta consciente e orientado, deambula sem auxílio, senta e levanta sem apoio, apresenta membros superiores e inferiores simétricos, com trofismo normal, força muscular preservada, sem sinais de hipoatividade ou hipofunção, ausência de lesões cicatriciais, sem evidências de hematomas ou equimoses recentes em face, tronco ou membros, equilíbrio preservado (Romberg negativo)."
    exame_clinico = st.text_area("Descreva o Exame Clínico:", value=exame_clinico_default if 'exame_clinico' not in st.session_state else st.session_state.exame_clinico, height=200)
    st.session_state.exame_clinico = exame_clinico

    col_exame1, col_exame2 = st.columns([1, 1])

    # --- NOVO BLOCO: Salvar Exame e Vincular Patologia por nome customizado ---
    with col_exame1:
        if "exames_salvos" not in st.session_state:
            st.session_state.exames_salvos = {}
        st.markdown("💾 Salvar Exame")
        nome_patologia_input = st.text_input("Informe o nome da patologia:", key="input_nome_patologia_salvar", help="", label_visibility="visible")
        if st.button("Salvar Exame Vinculado", key="btn_salvar_exame_patologia"):
            nome_pat = nome_patologia_input.strip()
            if nome_pat:
                st.session_state.exames_salvos[nome_pat] = exame_clinico
                st.success(f"Exame salvo e vinculado à patologia: {nome_pat}")
            else:
                st.warning("Informe o nome da patologia para salvar.")

    # --- NOVO BLOCO: Exames Padrão Salvos com busca, sugestão e exclusão ---
    with col_exame2:
        col_exames, col_trash = st.columns([8,1])
        with col_exames:
            st.markdown("📥 Exames Salvos")
        with col_trash:
            # Botão pequeno de lixeira para exclusão suspensa
            if st.button("🗑️", key="btn_trash_exames_salvos"):
                st.session_state.show_exames_delete = True
        # Excluir exame salvo
        if st.session_state.get("show_exames_delete", False):
            patologias_salvas = list(st.session_state.exames_salvos.keys())
            if len(patologias_salvas) == 0:
                st.info("Nenhuma patologia salva para excluir.")
            else:
                pat_del = st.selectbox("Selecione a patologia para excluir:", patologias_salvas, key="select_patologia_excluir")
                col_exc1, col_exc2 = st.columns(2)
                with col_exc1:
                    if st.button("Excluir", key="btn_excluir_patologia"):
                        del st.session_state.exames_salvos[pat_del]
                        st.session_state.show_exames_delete = False
                        st.rerun()
                with col_exc2:
                    if st.button("Cancelar", key="btn_cancelar_patologia_excluir"):
                        st.session_state.show_exames_delete = False
        else:
            # Busca e seleção de exames padrão salvos
            patologias_salvas = list(st.session_state.exames_salvos.keys())
            if len(patologias_salvas) == 0:
                st.info("Nenhum exame padrão foi salvo ainda.")
            else:
                busca_pat = st.text_input("Buscar patologia:", key="input_busca_patologia")
                sugestoes = [p for p in patologias_salvas if busca_pat.lower() in p.lower()] if busca_pat else patologias_salvas
                if not sugestoes:
                    st.info("Nenhuma patologia encontrada.")
                else:
                    pat_select = st.selectbox("Selecione uma patologia salva:", sugestoes, key="select_patologia_salva")
                    if st.button("Carregar Exame", key="btn_carregar_exame_patologia"):
                        st.session_state.exame_clinico = st.session_state.exames_salvos[pat_select]
                        st.success(f"Exame carregado para: {pat_select}")
                        st.rerun()

    st.subheader("🩻 Patologias")
    if "patologias" not in st.session_state:
        st.session_state.patologias = ["Transtorno de disco intervertebral cervical (CID M50)"]
    for idx, pat in enumerate(st.session_state.patologias):
        col1, col2 = st.columns([6,1])
        col1.markdown(f"- {pat}")
        if col2.button("❌", key=f"del_pat_{idx}"):
            st.session_state.patologias.pop(idx)
            st.rerun()

    patologias_memoria = sorted(lista_cids.items(), key=lambda x: x[0])
    lista_opcoes = [f"{v} (CID {k})" for k, v in patologias_memoria]
    nova_pat = st.selectbox("Selecione uma patologia para adicionar:", lista_opcoes)

    # NOVO BLOCO DE BOTÕES AO LADO DO GERENCIAR
    col_add, col_gerenciar = st.columns([2, 2])
    with col_add:
        if st.button("➕ Adicionar Patologia"):
            if nova_pat not in st.session_state.patologias:
                st.session_state.patologias.append(nova_pat)
                st.rerun()
    with col_gerenciar:
        if st.button("🛠️ Gerenciar"):
            st.session_state["abrir_gerenciador_patologias"] = True

    if st.session_state.get("abrir_gerenciador_patologias", False):
        st.markdown("### 🛠️ Gerenciar Patologias Salvas")
        patologias_ordenadas = sorted(lista_cids.items(), key=lambda x: x[0])
        lista_formatada = [f"{cid} - {desc}" for cid, desc in patologias_ordenadas]
        pat_selecionada = st.selectbox("Patologias na memória:", lista_formatada, key="gerenciar_lista")

        if st.button("➕ Adicionar Nova Patologia", key="btn_abrir_adicionar_patologia"):
            st.session_state["abrir_nova_patologia"] = True

        # Botão Editar Patologia da Memória
        if st.button("✏️ Editar Patologia da Memória", key="btn_editar_patologia_memoria"):
            st.session_state["editar_patologia_memoria"] = True

        if st.session_state.get("abrir_nova_patologia", False):
            novo_cid = st.text_input("Novo CID:", key="novo_cid_input")
            nova_desc = st.text_input("Descrição da patologia:", key="nova_desc_input")
            conceito = st.text_area("Conceito clínico:", key="conceito_input", height=150)
            # Novo bloco para salvar patologia
            if st.button("Salvar Patologia", key="btn_salvar_nova_patologia"):
                if novo_cid and nova_desc:
                    # Atualizar a lista_cids salva no session_state
                    if "patologias_memoria" not in st.session_state:
                        st.session_state["patologias_memoria"] = {}
                    st.session_state["patologias_memoria"][novo_cid] = {
                        "descricao": nova_desc,
                        "conceito": conceito
                    }

                    # Também atualizar a lista_cids diretamente
                    lista_cids[novo_cid] = nova_desc
                    if "conceitos" not in st.session_state:
                        st.session_state["conceitos"] = {}
                    st.session_state["conceitos"][novo_cid] = conceito
                    st.success(f"Patologia {novo_cid} - {nova_desc} incluída.")
                    st.session_state["abrir_nova_patologia"] = False
                    st.rerun()

        # Lógica condicional para editar patologia salva
        if st.session_state.get("editar_patologia_memoria", False):
            st.markdown("### ✏️ Editar Patologia Salva")
            cid_editar = st.selectbox("Selecione a patologia para editar:", lista_formatada, key="editar_lista")
            cid_atual, desc_atual = cid_editar.split(" - ", 1)
            novo_cid_editar = st.text_input("Novo CID:", value=cid_atual, key="cid_editar_input")
            nova_desc_editar = st.text_input("Nova descrição:", value=desc_atual, key="desc_editar_input")
            conceito_editar = st.text_area("Novo conceito clínico:", value=st.session_state.get("conceitos", {}).get(cid_atual, ""), key="conceito_editar_input")

            if st.button("Salvar Edição", key="btn_salvar_edicao"):
                # Remove o antigo
                lista_cids.pop(cid_atual, None)
                conceito_antigo = st.session_state.get("conceitos", {}).pop(cid_atual, None)

                # Adiciona o novo
                lista_cids[novo_cid_editar] = nova_desc_editar
                if "conceitos" not in st.session_state:
                    st.session_state["conceitos"] = {}
                st.session_state["conceitos"][novo_cid_editar] = conceito_editar

                st.success(f"Patologia {cid_atual} atualizada para {novo_cid_editar}.")
                st.session_state["editar_patologia_memoria"] = False
                st.rerun()

        if st.button("Fechar", key="btn_fechar_gerenciador"):
            st.session_state["abrir_gerenciador_patologias"] = False
            st.rerun()

    st.subheader("📊 Análise Pericial")
    decisao = st.radio("Selecione a decisão pericial:", ["Conceder", "Negar"], horizontal=True)

    if "justificativas_memoria" not in st.session_state:
        st.session_state.justificativas_memoria = {
            "Conceder": ["Crises frequentes", "Exame clínico compatível"],
            "Negar": ["Sem evidências clínicas", "Outros"]
        }

    just_memoria = st.session_state.justificativas_memoria[decisao]
    just_memoria.sort()

    excluir_key = f"excluir_{decisao.lower()}"
    if excluir_key not in st.session_state:
        st.session_state[excluir_key] = False

    st.markdown(f"**Justificativas para {decisao}:**")
    justificativas_selecionadas = []
    cols = st.columns(len(just_memoria) if just_memoria else 1)
    for idx, justificativa in enumerate(just_memoria):
        with cols[idx]:
            if st.checkbox(justificativa, key=f"{decisao}_{justificativa}"):
                justificativas_selecionadas.append(justificativa)
            if st.session_state[excluir_key]:
                if st.button("❌", key=f"del_{decisao}_{justificativa}"):
                    just_memoria.remove(justificativa)
                    just_memoria.sort()
                    st.rerun()

    nova_justificativa = st.text_input(f"Nova justificativa para {decisao}", key=f"input_nova_justificativa_{decisao}")
    col_add_just, col_del_just = st.columns([1,1])
    with col_add_just:
        if st.button("➕ Adicionar", key=f"btn_add_justificativa_{decisao}"):
            if nova_justificativa and nova_justificativa not in just_memoria:
                just_memoria.append(nova_justificativa)
                just_memoria.sort()
                st.rerun()
    with col_del_just:
        label = "❌ Excluir Ativado" if st.session_state[excluir_key] else "❌ Excluir"
        if st.button(label, key=f"btn_toggle_excluir_{decisao}"):
            st.session_state[excluir_key] = not st.session_state[excluir_key]

    if "Outros" in justificativas_selecionadas:
        st.text_area("Descreva outras justificativas")

    st.subheader("✅ Conclusão")
    conclusao_opcao = st.radio("Conclusão:", ["Não reconheço incapacidade", "Houve incapacidade apenas pretérita", "Reconheço a incapacidade atual"])
    st.session_state["conclusao_opcao"] = conclusao_opcao

    if conclusao_opcao == "Houve incapacidade apenas pretérita":
        data_ini = st.date_input("Data de início", format="DD/MM/YYYY", key="data_inicio_pret")
        data_fim = st.date_input("Data de fim", format="DD/MM/YYYY", key="data_fim_pret")
        reducao = st.checkbox("Há redução de capacidade?")
        if reducao:
            reducao_data = st.date_input("Data de início da redução de capacidade", format="DD/MM/YYYY", key="reducao_data")

    elif conclusao_opcao == "Reconheço a incapacidade atual":
        tipo = st.selectbox("Tipo:", ["Omniprofissional", "Multiprofissional", "Uniprofissional"], key="tipo_incapacidade")
        permanente = st.radio("É permanente?", ["Sim", "Não"], key="permanente")
        if permanente == "Não":
            duracao = st.number_input("Duração estimada (meses)", min_value=1, max_value=60, key="duracao_incap")
        if tipo == "Multiprofissional":
            profissao_multiprof = st.text_input("Profissão atual do autor:", value=profissao, key="prof_multiprof")
            prof_restritas = st.text_area("Quais atividades não pode exercer?", key="ativ_restritas")
        elif tipo == "Uniprofissional":
            profissao_uniprof = st.text_input("Profissão atual do autor:", value=profissao, key="prof_uniprof")
        dii = st.date_input("Data de início da incapacidade (DII)", format="DD/MM/YYYY", key="dii_data")
        terceiro = st.checkbox("Necessita de terceiro?", key="necessita_terceiro")

    elif conclusao_opcao == "Não reconheço incapacidade":
        reducao = st.checkbox("Há redução de capacidade?", key="reducao_nao")
        if reducao:
            reducao_data = st.date_input("Data de início da redução de capacidade", format="DD/MM/YYYY", key="reducao_data_nao")

    # --- Redação automática da conclusão ---
    conclusao_texto = "____________________________"
    if conclusao_opcao == "Não reconheço incapacidade":
        if st.session_state.get("dados_extraidos", {}).get("der") and st.session_state.get("historico_beneficios", "") != "Nega benefício previdenciário prévio.":
            conclusao_texto = (
                "Dessa forma, diante da anamnese, exame clínico pericial e documentos apresentados, "
                "reconheço como superada incapacidade laboral da parte autora, sem evidências de alterações "
                "que fundamentem persistência da incapacidade."
            )
        else:
            conclusao_texto = (
                "Dessa forma, diante da anamnese, exame clínico pericial e documentos apresentados, "
                "não reconheço incapacidade laboral da parte autora."
            )
    elif conclusao_opcao == "Houve incapacidade apenas pretérita":
        dt_ini = st.session_state.get("data_inicio_pret")
        dt_fim = st.session_state.get("data_fim_pret")
        texto_periodo = f"{dt_ini.strftime('%d/%m/%Y')} à {dt_fim.strftime('%d/%m/%Y')}" if dt_ini and dt_fim else "período não informado"
        conclusao_texto = (
            f"Desta forma, diante da anamnese, exame físico pericial e documentos analisados, "
            f"reconheço que houve incapacidade laborativa omniprofissional no período {texto_periodo}. "
            "Ademais, superada a incapacidade, sem sinais de agravamento da patologia que sugiram persistência de quadro incapacitante."
        )
        if st.session_state.get("reducao_pret"):
            data_reducao = st.session_state.get("reducao_data")
            if data_reducao:
                conclusao_texto += (
                    f"\nPor conseguinte, em decorrência das sequelas que acometem a parte autora, reconheço a redução da sua capacidade laboral, "
                    f"atribuindo como marco temporal inicial a data de {data_reducao.strftime('%d/%m/%Y')}."
                )
    elif conclusao_opcao == "Reconheço a incapacidade atual":
        tipo = st.session_state.get("tipo_incapacidade")
        dii = st.session_state.get("dii_data")
        if st.session_state.get("permanente") == "Sim":
            conclusao_texto = (
                f"Desta forma, diante da anamnese, exame físico pericial e documentos analisados, "
                f"reconheço a incapacidade laborativa {tipo.lower() if tipo else ''}, tendo como DII a data de {dii.strftime('%d/%m/%Y') if dii else '___'}. "
                "Ademais, diante da cronicidade e avançar do quadro clínico, reconheço a incapacidade como permanente."
            )
        else:
            duracao = st.session_state.get("duracao_incap")
            conclusao_texto = (
                f"Desta forma, diante da anamnese, exame físico pericial e documentos analisados, "
                f"reconheço a incapacidade laborativa {tipo.lower() if tipo else ''}, tendo como DII a data de {dii.strftime('%d/%m/%Y') if dii else '___'}. "
                f"Ademais, diante do quadro clínico, estimo necessário afastamento das atividades pelo período de {duracao} meses a contar desta perícia."
            )
    st.session_state["conclusao_gerada"] = conclusao_texto


# Fim dos campos do laudo principais
# Adicione os novos botões após o bloco de conclusão:
    col_botoes_finais = st.columns([1, 1, 1])
    with col_botoes_finais[0]:
        if st.button("💾 Salvar"):
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
            doc = Document()

            style = doc.styles["Normal"]
            font = style.font
            font.name = "Palatino Linotype"
            font.size = Pt(12)

            def add_paragrafo(texto, bold=False, underline=False, align="JUSTIFY", space_before=0, space_after=0):
                p = doc.add_paragraph()
                run = p.add_run(texto)
                run.bold = bold
                run.underline = underline
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                if align == "CENTER":
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif align == "RIGHT":
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            add_paragrafo("PROCESSO DE AUTOS Nº: 0004137-16.2025.4.05.8102".upper(), space_after=0)
            add_paragrafo(f"AUTOR(A): {nome_autor}".upper(), space_after=0)
            add_paragrafo("RÉU: INSTITUTO NACIONAL DO SEGURO SOCIAL – INSS", space_after=0)
            add_paragrafo("PERITO(A) NOMEADO(A): CICERO HYTTALLO CARNEIRO BALDUINO", space_after=0)
            add_paragrafo("Nº DE INSCRIÇÃO NO CRM/UF: CRM-CE nº 21.652", space_after=0)
            add_paragrafo("LAUDO MÉDICO PERICIAL", align="CENTER", space_after=0)
            if "tipo_beneficio" not in st.session_state:
                st.session_state["tipo_beneficio"] = "AD"

            tipo_laudo = st.session_state["tipo_beneficio"]
            if tipo_laudo == "AD":
                subtitulo = "AUXÍLIO PREVIDENCIÁRIO POR INCAPACIDADE TEMPORÁRIA"
            elif tipo_laudo == "BPC":
                subtitulo = "AUXÍLIO ASSISTENCIAL AO IDOSO OU DEFICIENTE - BPC"
            elif tipo_laudo == "DPVAT":
                subtitulo = "SEGURO DPVAT"
            elif tipo_laudo == "INT":
                subtitulo = "INTERDIÇÃO"
            else:
                subtitulo = "TIPO DE LAUDO NÃO ESPECIFICADO"

            add_paragrafo(subtitulo, bold=True, align="CENTER", space_after=0)

            meses_pt = {
                "January": "Janeiro", "February": "Fevereiro", "March": "Março", "April": "Abril",
                "May": "Maio", "June": "Junho", "July": "Julho", "August": "Agosto",
                "September": "Setembro", "October": "Outubro", "November": "Novembro", "December": "Dezembro"
            }
            data_hoje = datetime.today()
            mes_extenso = meses_pt[data_hoje.strftime("%B").capitalize()]
            data_formatada = f"{data_hoje.strftime('%d')} de {mes_extenso} de {data_hoje.strftime('%Y')}"
            add_paragrafo(f"Data de realização do exame: {data_formatada}", align="CENTER", space_after=0)
            # Linha removida para garantir que "APRESENTAÇÃO" venha imediatamente após o subtítulo

            add_paragrafo("APRESENTAÇÃO", bold=True, align="CENTER", space_after=0)
            add_paragrafo("Periciando(a): " + nome_autor, space_after=0)
            add_paragrafo("Data de nascimento: " + data_nascimento, space_after=0)
            add_paragrafo("Idade: ______ anos", space_after=0)
            add_paragrafo("RG: " + st.session_state.get("dados_extraidos", {}).get('rg', ''), space_after=0)
            add_paragrafo("CPF: " + st.session_state.get("dados_extraidos", {}).get('cpf', ''), space_after=0)
            add_paragrafo("Profissão referida: " + profissao, space_after=0)
            add_paragrafo("Histórico laboral prévio: _______________________", space_after=0)
            add_paragrafo("Escolaridade: _______________________", space_after=0)
            add_paragrafo("DER: " + st.session_state.get("dados_extraidos", {}).get('der', ''), space_after=0)
            add_paragrafo("", space_after=0)

            add_paragrafo("ANAMNESE", bold=True, align="CENTER", space_after=0)
            for paragrafo in anamnese.split("\n"):
                if paragrafo.strip():
                    add_paragrafo(paragrafo.strip(), space_after=0)
            add_paragrafo("", space_after=0)

            add_paragrafo("EXAME CLÍNICO", bold=True, align="CENTER", space_after=0)
            for paragrafo in exame_clinico.split("\n"):
                if paragrafo.strip():
                    add_paragrafo(paragrafo.strip(), space_after=0)
            add_paragrafo("", space_after=0)

            # Os campos de documentos, benefícios, fundamentação, análise e conclusão não estão mais aqui, pois removidos.

            # Salvar o laudo gerado
            nome_arquivo = f"Laudo_{nome_autor.replace(' ', '_')}_{datetime.today().strftime('%Y%m%d')}.docx"
            caminho_arquivo = os.path.join(os.getcwd(), nome_arquivo)
            doc.save(caminho_arquivo)

            # Download automático se sinalizado
            if st.session_state.get("gerar_pdf_automaticamente", False):
                st.session_state["gerar_pdf_automaticamente"] = False
                st.success("Download automático iniciado.")
                with open(caminho_arquivo, "rb") as f:
                    base64_pdf = base64.b64encode(f.read()).decode('utf-8')
                    href = f'<meta http-equiv="refresh" content="0;url=data:application/octet-stream;base64,{base64_pdf}" download="{nome_arquivo}">'
                    st.markdown(href, unsafe_allow_html=True)

            st.success(f"Laudo salvo com sucesso como: {nome_arquivo}")
            with open(caminho_arquivo, "rb") as f:
                st.download_button("📄 Baixar Laudo Final", f, file_name=nome_arquivo)

    with col_botoes_finais[1]:
        if st.button("🧠 Responder Quesitos"):
            st.info("Respostas aos quesitos geradas com base no conteúdo do laudo.")  # Lógica a ser integrada

    with col_botoes_finais[2]:
        if st.button("📥 Finalizar"):
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            doc = Document()
            style = doc.styles["Normal"]
            font = style.font
            font.name = "Palatino Linotype"
            font.size = Pt(12)

            def add_paragrafo(texto, bold=False, underline=False, align="JUSTIFY", space_before=0, space_after=0):
                p = doc.add_paragraph()
                run = p.add_run(texto)
                run.bold = bold
                run.underline = underline
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                if align == "CENTER":
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif align == "RIGHT":
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Cabeçalho em maiúsculo
            # Número do processo
            numero_processo_match = re.search(r"N[uú]mero[:\s]*([0-9]{7}-[0-9]{2}\.[0-9]{4}\.[0-9]\.[0-9]{2}\.[0-9]{4})", texto_extraido, re.IGNORECASE) if 'texto_extraido' in locals() or 'texto_extraido' in globals() else None
            numero_processo = numero_processo_match.group(1) if numero_processo_match else "NÃO ENCONTRADO"
            nome_autor = st.session_state.get("dados_extraidos", {}).get("nome", "")
            nome_autor_maiusculo = nome_autor.upper() if nome_autor else ""
            data_nasc = st.session_state.get("dados_extraidos", {}).get("nascimento", "")
            idade = calcular_idade(data_nasc) if data_nasc else ""
            rg = st.session_state.get("dados_extraidos", {}).get("rg", "")
            cpf = st.session_state.get("dados_extraidos", {}).get("cpf", "")
            profissao = st.session_state.get("profissao", "") or st.session_state.get("dados_extraidos", {}).get("profissao", "")
            der = st.session_state.get("dados_extraidos", {}).get("der", "")
            historico_laboral = st.session_state.get("historico_laboral", "")
            escolaridade = []
            if st.session_state.get("esc_analfabeto"): escolaridade.append("Analfabeto")
            if st.session_state.get("esc_so_assina"): escolaridade.append("Só assina o nome")
            if st.session_state.get("esc_nao_sabe"): escolaridade.append("Não sabe informar")
            if st.session_state.get("esc_fund_inc"): escolaridade.append("Fundamental incompleto")
            if st.session_state.get("esc_fund_comp"): escolaridade.append("Fundamental completo")
            if st.session_state.get("esc_med_inc"): escolaridade.append("Médio incompleto")
            if st.session_state.get("esc_med_comp"): escolaridade.append("Médio completo")
            if st.session_state.get("esc_sup_inc"): escolaridade.append("Superior incompleto")
            if st.session_state.get("esc_sup_comp"): escolaridade.append("Superior completo")
            escolaridade_str = ", ".join(escolaridade) if escolaridade else "Não informada"
            anamnese = st.session_state.get("anamnese", "")
            exame_clinico = st.session_state.get("exame_clinico", "")
            documentos_medicos = st.session_state.get("documentos_medicos", "") if "documentos_medicos" in st.session_state else "____________________________"
            historico_beneficios = "Nega benefício previdenciário prévio."
            # Novo bloco: FUNDAMENTAÇÃO TEÓRICA conforme instrução
            fundamentacao_teorica = ""
            patologias = st.session_state.get("patologias", [])
            conceitos = st.session_state.get("conceitos", {})
            # Extrair nomes e CIDs
            pat_infos = []
            for pat in patologias:
                cid_match = re.search(r"CID\s*([A-Z0-9\.]+)", pat)
                nome_pat = pat
                cid = None
                if cid_match:
                    cid = cid_match.group(1)
                    nome_pat = pat[:pat.find("(CID")].strip() if "(CID" in pat else pat.strip()
                else:
                    nome_pat = pat.strip()
                if cid:
                    pat_infos.append((nome_pat, cid))
            if pat_infos:
                # Montar frase introdutória
                intro = "Com base na petição inicial e documentos médicos, extrai-se que a parte autora é portadora de "
                pats_txt = []
                for i, (nome, cid) in enumerate(pat_infos):
                    if i == len(pat_infos) - 1 and i > 0:
                        pats_txt.append(f"e de {nome} (CID {cid})")
                    else:
                        pats_txt.append(f"{nome} (CID {cid})")
                if len(pats_txt) > 2:
                    intro += ", ".join(pats_txt[:-1]) + f", {pats_txt[-1]}."
                elif len(pats_txt) == 2:
                    intro += f"{pats_txt[0]} {pats_txt[1]}."
                else:
                    intro += f"{pats_txt[0]}."
                fundamentacao_teorica += intro + "\n"
                # Inserir cada conceito por CID (ordem listada)
                for nome, cid in pat_infos:
                    conceito = conceitos.get(cid, "")
                    if conceito:
                        fundamentacao_teorica += f"{conceito}\n\n"
            else:
                fundamentacao_teorica = "____________________________"
            decisao = st.session_state.get("decisao", "")
            justificativas = []
            for key in ["Conceder", "Negar"]:
                if key == decisao:
                    justs = st.session_state.get("justificativas_memoria", {}).get(key, [])
                    for j in justs:
                        if st.session_state.get(f"{key}_{j}", False):
                            justificativas.append(j)
            import openai  # Certifique-se de que a API OpenAI está configurada em seu ambiente
            analise_pericial = ""
            if justificativas:
                prompt_analise = (
                    f"Redija um parágrafo em linguagem técnica médica e pericial com base nas justificativas selecionadas a seguir. "
                    f"O contexto é uma perícia judicial de auxílio-doença. Justificativas: {', '.join(justificativas)}."
                )
                try:
                    resposta = openai.ChatCompletion.create(
                        model="gpt-4",
                        messages=[
                            {"role": "system", "content": "Você é um médico perito judicial que redige análises periciais formais com base em justificativas clínicas."},
                            {"role": "user", "content": prompt_analise}
                        ],
                        max_tokens=400
                    )
                    analise_pericial = resposta["choices"][0]["message"]["content"].strip()
                except Exception as e:
                    analise_pericial = "Erro ao gerar análise pericial automaticamente. Justificativas selecionadas: " + "; ".join(justificativas)
            else:
                analise_pericial = "____________________________"
            # Usar o texto gerado automaticamente na conclusão
            conclusao_texto = st.session_state.get("conclusao_gerada", "____________________________")
            # Quesitos
            quesitos_juiz = "Resposta: ____________________________"
            quesitos_autor = st.session_state.get("quesitos_autor", "")
            quesitos_reu = "Resposta: ____________________________"

            # Títulos centrais
            add_paragrafo(f"PROCESSO DE AUTOS Nº: {numero_processo.upper()}", bold=False, align="JUSTIFY", space_after=0)
            add_paragrafo(f"AUTOR(A): {nome_autor_maiusculo}", bold=False, align="JUSTIFY", space_after=0)
            add_paragrafo("RÉU: INSTITUTO NACIONAL DO SEGURO SOCIAL – INSS", bold=False, align="JUSTIFY", space_after=0)
            add_paragrafo("PERITO(A) NOMEADO(A): CICERO HYTTALLO CARNEIRO BALDUINO", bold=False, align="JUSTIFY", space_after=0)
            add_paragrafo("Nº DE INSCRIÇÃO NO CRM/UF: CRM-CE nº 21.652", bold=False, align="JUSTIFY", space_after=0)
            add_paragrafo("", align="JUSTIFY", space_after=0)
            add_paragrafo("LAUDO MÉDICO PERICIAL", bold=True, align="CENTER", space_after=0)

            # Novo bloco: subtítulo conforme tipo_beneficio
            if "tipo_beneficio" not in st.session_state:
                st.session_state["tipo_beneficio"] = "AD"

            tipo_laudo = st.session_state["tipo_beneficio"]
            if tipo_laudo == "AD":
                subtitulo = "Auxílio previdenciário por Incapacidade Temporária"
            elif tipo_laudo == "BPC":
                subtitulo = "Auxílio assistencial ao Idoso ou Deficiente - BPC"
            elif tipo_laudo == "DPVAT":
                subtitulo = "Seguro DPVAT"
            elif tipo_laudo == "INT":
                subtitulo = "Interdição"
            else:
                subtitulo = "Tipo de laudo não especificado"

            add_paragrafo(subtitulo.upper(), bold=True, align="CENTER", space_after=0)

            # Data do exame (em português)
            meses_pt = {
                "January": "janeiro", "February": "fevereiro", "March": "março", "April": "abril",
                "May": "maio", "June": "junho", "July": "julho", "August": "agosto",
                "September": "setembro", "October": "outubro", "November": "novembro", "December": "dezembro"
            }
            data_hoje = datetime.today()
            mes_extenso = meses_pt[data_hoje.strftime("%B")]
            data_formatada = f"{data_hoje.day:02d} de {mes_extenso} de {data_hoje.year}"
            # Troca bloco data e apresentação conforme solicitado
            add_paragrafo("", align="JUSTIFY", space_after=0)
            add_paragrafo("APRESENTAÇÃO", bold=True, align="CENTER", space_after=0)
            add_paragrafo(f"Data da realização do exame: {data_formatada}", align="JUSTIFY", space_after=0)
            add_paragrafo("Local da perícia: 17ª Vara Federal", align="JUSTIFY", space_after=0)
            add_paragrafo("Perito: Cicero Hyttallo Carneiro Balduino; CRM/CE nº 21.652", align="JUSTIFY", space_after=0)
            add_paragrafo("Especialidade do perito: Médico Pós-graduado em Medicina do Trabalho; Pós-graduado em Auditoria em Serviços de Saúde; Pós-graduado em Perícia Forense; Médico Auditor; Graduado em Direito.", align="JUSTIFY", space_after=0)
            add_paragrafo("Assistente técnico da parte autora: ---", align="JUSTIFY", space_after=0)
            add_paragrafo("Assistente técnico da parte ré: ---", align="JUSTIFY", space_after=0)
            add_paragrafo("", align="JUSTIFY", space_after=0)
            add_paragrafo("IDENTIFICAÇÃO DO(A) PERICIANDO(A)", bold=True, align="CENTER", space_after=0)

            # Bloco de identificação do periciando(a)
            add_paragrafo(f"Periciando(a):\t{nome_autor}", align="LEFT", space_after=0)
            add_paragrafo(f"Data de nascimento:\t{data_nasc}", align="LEFT", space_after=0)
            add_paragrafo(f"Idade: {idade}", align="LEFT", space_after=0)
            add_paragrafo(f"RG: {rg}", align="LEFT", space_after=0)
            add_paragrafo(f"CPF: {cpf}", align="LEFT", space_after=0)
            add_paragrafo(f"Profissão referida: {profissao}", align="LEFT", space_after=0)
            add_paragrafo(f"Histórico laboral prévio: {historico_laboral}", align="LEFT", space_after=0)
            add_paragrafo(f"Escolaridade: {escolaridade_str}", align="LEFT", space_after=0)
            add_paragrafo(f"DER: {der}", align="LEFT", space_after=0)

            # ANAMNESE
            add_paragrafo("ANAMNESE", bold=True, align="CENTER", space_after=0)
            for par in anamnese.split("\n"):
                if par.strip():
                    add_paragrafo(par, align="JUSTIFY", space_after=0)
            add_paragrafo("", align="JUSTIFY", space_after=0)

            # EXAME CLÍNICO
            add_paragrafo("EXAME CLÍNICO", bold=True, align="CENTER", space_after=0)
            for par in exame_clinico.split("\n"):
                if par.strip():
                    add_paragrafo(par, align="JUSTIFY", space_after=0)
            add_paragrafo("", align="JUSTIFY", space_after=0)

            # DOCUMENTOS MÉDICOS
            add_paragrafo("DOCUMENTOS MÉDICOS", bold=True, align="CENTER", space_after=0)
            add_paragrafo("Documentos médicos ao final anexados.", align="JUSTIFY", space_after=0)
            add_paragrafo("", align="JUSTIFY", space_after=0)

            # HISTÓRICO DE BENEFÍCIOS
            add_paragrafo("HISTÓRICO DE BENEFÍCIOS", bold=True, align="CENTER", space_after=0)
            add_paragrafo(historico_beneficios, align="JUSTIFY", space_after=0)
            add_paragrafo("", align="JUSTIFY", space_after=0)

            # FUNDAMENTAÇÃO TEÓRICA
            add_paragrafo("FUNDAMENTAÇÃO TEÓRICA", bold=True, align="CENTER", space_after=0)
            for par in fundamentacao_teorica.split("\n"):
                if par.strip():
                    add_paragrafo(par, align="JUSTIFY", space_after=0)
            add_paragrafo("", align="JUSTIFY", space_after=0)

            # ANÁLISE PERICIAL
            add_paragrafo("ANÁLISE PERICIAL", bold=True, align="CENTER", space_after=0)
            add_paragrafo(analise_pericial, align="JUSTIFY", space_after=0)
            add_paragrafo("", align="JUSTIFY", space_after=0)

            # CONCLUSÃO
            add_paragrafo("CONCLUSÃO", bold=True, align="CENTER", space_after=0)
            add_paragrafo(conclusao_texto, align="JUSTIFY", space_after=0)
            add_paragrafo("", align="JUSTIFY", space_after=0)

            # QUESITOS DO JUIZ
            add_paragrafo("QUESITOS DO JUIZ", bold=True, align="CENTER", space_after=0)
            add_paragrafo(quesitos_juiz, align="JUSTIFY", space_after=0)
            add_paragrafo("", align="JUSTIFY", space_after=0)
            # QUESITOS DO AUTOR
            if st.session_state.get("tem_quesitos_autor", False):
                add_paragrafo("QUESITOS DO(A) AUTOR(A)", bold=True, align="CENTER", space_after=0)
                add_paragrafo(quesitos_autor if quesitos_autor else "____________________________", align="JUSTIFY", space_after=0)
                add_paragrafo("", align="JUSTIFY", space_after=0)
            # QUESITOS DO RÉU
            add_paragrafo("QUESITOS DO RÉU", bold=True, align="CENTER", space_after=0)
            add_paragrafo(quesitos_reu, align="JUSTIFY", space_after=0)
            add_paragrafo("", align="JUSTIFY", space_after=0)

            # Assinatura
            add_paragrafo("CICERO HYTTALLO CARNEIRO BALDUINO", bold=True, align="CENTER", space_after=0)
            add_paragrafo("Médico Perito Judicial - CRM/CE nº 21.652", align="CENTER", space_after=0)

            # Quebra de página para anexos
            doc.add_page_break()
            add_paragrafo("ANEXOS", bold=True, align="CENTER", space_after=0)
            doc.add_paragraph("\n\n")
            imagens = st.session_state.get("anexos_docs", [])
            if imagens:
                for imagem in imagens:
                    doc.add_picture(imagem, width=Inches(5))
                    doc.add_page_break()

            # Salvar e baixar
            nome_arquivo_final = f"Laudo_Final_{datetime.today().strftime('%Y%m%d')}.docx"
            caminho_arquivo_final = os.path.join(os.getcwd(), nome_arquivo_final)
            doc.save(caminho_arquivo_final)
            st.success("Laudo final gerado com sucesso.")
            with open(caminho_arquivo_final, "rb") as f:
                st.download_button("📄 Baixar Laudo com Anexos", f, file_name=nome_arquivo_final)

    # Exibição de quesitos do autor extraídos, se houver
    if st.session_state.get("tem_quesitos_autor", False):
        st.markdown("### 📄 Quesitos do Autor")
        st.text_area("Quesitos extraídos do processo para análise:", value=st.session_state.get("quesitos_autor", ""), height=200)

    # Novo bloco: geração do laudo final com anexos
    if st.session_state.get("finalizar_laudo", False):
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        doc = Document()

        # Estilo
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Palatino Linotype"
        font.size = Pt(12)

        def add_paragrafo(texto, bold=False, underline=False, align="JUSTIFY", space_before=0, space_after=0):
            p = doc.add_paragraph()
            run = p.add_run(texto)
            run.bold = bold
            run.underline = underline
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if align == "CENTER":
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif align == "RIGHT":
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        add_paragrafo("CICERO HYTTALLO CARNEIRO BALDUINO", bold=True, align="CENTER", space_after=0)
        add_paragrafo("Médico Perito Judicial - CRM/CE nº 21.652", align="CENTER", space_after=0)

        # Nova página
        doc.add_page_break()
        add_paragrafo("ANEXOS", bold=True, align="CENTER", space_after=0)
        doc.add_paragraph("\n\n")  # Pula duas linhas

        imagens = st.session_state.get("anexos_docs", [])
        for imagem in imagens:
            doc.add_picture(imagem, width=Inches(5))
            doc.add_page_break()

        nome_arquivo_final = f"Laudo_Final_{datetime.today().strftime('%Y%m%d')}.docx"
        caminho_arquivo_final = os.path.join(os.getcwd(), nome_arquivo_final)
        doc.save(caminho_arquivo_final)

        st.success("Laudo final com anexos gerado.")
        with open(caminho_arquivo_final, "rb") as f:
            st.download_button("📄 Baixar Laudo com Anexos", f, file_name=nome_arquivo_final)

        st.session_state["finalizar_laudo"] = False

# Remover chamada direta se existir e garantir execução correta
# Se desejar rodar com python laudos_ad.py, descomente abaixo:
if __name__ == "__main__":
    interface_laudo_ad()
