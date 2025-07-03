import streamlit as st
import pandas as pd
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from PIL import Image
import base64

# Configuração da página
st.set_page_config(
    page_title="Gerador de Laudos Periciais",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS personalizado
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        text-align: center;
    }
    .main-header h1 {
        color: white;
        font-size: 2.5rem;
        margin-bottom: 0.5rem;
    }
    .main-header p {
        color: #e8f4f8;
        font-size: 1.1rem;
        margin: 0;
    }
    .section-header {
        background: #f8f9fa;
        padding: 1rem;
        border-left: 4px solid #2a5298;
        margin: 1rem 0;
        border-radius: 5px;
    }
    .info-box {
        background: #e7f3ff;
        border: 1px solid #b3d9ff;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .success-box {
        background: #d4edda;
        border: 1px solid #c3e6cb;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .stButton > button {
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        border: none;
        padding: 0.75rem 2rem;
        border-radius: 5px;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.2);
    }
    .upload-section {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border: 2px dashed #dee2e6;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Função para criar cabeçalho da aplicação
def create_header():
    st.markdown("""
    <div class="main-header">
        <h1>⚖️ Gerador de Laudos Periciais</h1>
        <p>Sistema Inteligente para Elaboração de Perícias Médicas</p>
    </div>
    """, unsafe_allow_html=True)

# Função para validar CPF
def validar_cpf(cpf):
    cpf = ''.join(filter(str.isdigit, cpf))
    if len(cpf) != 11:
        return False
    
    if cpf == cpf[0] * 11:
        return False
    
    # Validação do primeiro dígito
    soma = sum(int(cpf[i]) * (10 - i) for i in range(9))
    resto = soma % 11
    if resto < 2:
        digito1 = 0
    else:
        digito1 = 11 - resto
    
    if int(cpf[9]) != digito1:
        return False
    
    # Validação do segundo dígito
    soma = sum(int(cpf[i]) * (11 - i) for i in range(10))
    resto = soma % 11
    if resto < 2:
        digito2 = 0
    else:
        digito2 = 11 - resto
    
    return int(cpf[10]) == digito2

# Função para formatar CPF
def formatar_cpf(cpf):
    cpf = ''.join(filter(str.isdigit, cpf))
    if len(cpf) == 11:
        return f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}"
    return cpf
    # Função para configurar cabeçalho do documento
def configurar_cabecalho(doc):
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "LAUDO PERICIAL MÉDICO"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Configurar fonte do cabeçalho
    run = header_para.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True

# Função para configurar rodapé do documento
def configurar_rodape(doc):
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Página "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar número da página
    run = footer_para.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Adicionar campo de número de página
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)

# Função para adicionar quebra de página
def adicionar_quebra_pagina(doc):
    doc.add_page_break()

# Função para adicionar título centralizado
def adicionar_titulo(doc, texto, tamanho=14, negrito=True):
    para = doc.add_paragraph()
    run = para.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(tamanho)
    run.bold = negrito
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para

# Função para adicionar parágrafo formatado
def adicionar_paragrafo(doc, texto, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, tamanho=12):
    para = doc.add_paragraph()
    run = para.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(tamanho)
    para.alignment = alinhamento
    return para

# Função para adicionar seção com título
def adicionar_secao(doc, titulo, conteudo, tamanho_titulo=12):
    # Título da seção
    para_titulo = doc.add_paragraph()
    run_titulo = para_titulo.add_run(titulo)
    run_titulo.font.name = 'Times New Roman'
    run_titulo.font.size = Pt(tamanho_titulo)
    run_titulo.bold = True
    para_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Conteúdo da seção
    para_conteudo = doc.add_paragraph()
    run_conteudo = para_conteudo.add_run(conteudo)
    run_conteudo.font.name = 'Times New Roman'
    run_conteudo.font.size = Pt(12)
    para_conteudo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    return para_titulo, para_conteudo

# Função para processar e redimensionar imagem
def processar_imagem(imagem_bytes, largura_maxima=3, altura_maxima=4):
    """
    Processa a imagem para o tamanho adequado (3x4 para foto do autor)
    """
    try:
        img = Image.open(io.BytesIO(imagem_bytes))
        
        # Converter para RGB se necessário
        if img.mode in ('RGBA', 'LA', 'P'):
            img = img.convert('RGB')
        
        # Redimensionar mantendo proporção
        img.thumbnail((int(largura_maxima * 100), int(altura_maxima * 100)), Image.Resampling.LANCZOS)
        
        # Salvar em bytes
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='JPEG', quality=85)
        img_bytes.seek(0)
        
        return img_bytes
    except Exception as e:
        st.error(f"Erro ao processar imagem: {str(e)}")
        return None

# Função para adicionar imagem ao documento
def adicionar_imagem_documento(doc, imagem_bytes, largura_inches=2.0, altura_inches=2.7):
    """
    Adiciona imagem ao documento Word
    """
    try:
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(imagem_bytes, width=Inches(largura_inches), height=Inches(altura_inches))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    except Exception as e:
        st.error(f"Erro ao adicionar imagem ao documento: {str(e)}")
        return False
        def main():
    create_header()
    
    # Sidebar para navegação
    st.sidebar.title("📋 Navegação")
    opcao = st.sidebar.selectbox(
        "Selecione a seção:",
        ["Dados do Processo", "Dados do Autor", "Dados da Perícia", "Anexos", "Gerar Laudo"]
    )
    
    # Inicializar session_state
    if 'dados_processo' not in st.session_state:
        st.session_state.dados_processo = {}
    if 'dados_autor' not in st.session_state:
        st.session_state.dados_autor = {}
    if 'dados_pericia' not in st.session_state:
        st.session_state.dados_pericia = {}
    if 'foto_autor' not in st.session_state:
        st.session_state.foto_autor = None
    if 'fotos_documentos' not in st.session_state:
        st.session_state.fotos_documentos = []
    
    # Seção: Dados do Processo
    if opcao == "Dados do Processo":
        st.markdown('<div class="section-header"><h2>📁 Dados do Processo</h2></div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.session_state.dados_processo['numero_processo'] = st.text_input(
                "Número do Processo:",
                value=st.session_state.dados_processo.get('numero_processo', ''),
                help="Digite o número completo do processo"
            )
            
            st.session_state.dados_processo['vara'] = st.text_input(
                "Vara/Juízo:",
                value=st.session_state.dados_processo.get('vara', ''),
                help="Ex: 1ª Vara Cível da Comarca de..."
            )
            
            st.session_state.dados_processo['comarca'] = st.text_input(
                "Comarca:",
                value=st.session_state.dados_processo.get('comarca', ''),
                help="Nome da comarca"
            )
        
        with col2:
            st.session_state.dados_processo['juiz'] = st.text_input(
                "Juiz(a):",
                value=st.session_state.dados_processo.get('juiz', ''),
                help="Nome do magistrado responsável"
            )
            
            st.session_state.dados_processo['tipo_acao'] = st.selectbox(
                "Tipo de Ação:",
                ["", "Previdenciária", "Securitária", "Trabalhista", "Cível", "Outras"],
                index=["", "Previdenciária", "Securitária", "Trabalhista", "Cível", "Outras"].index(
                    st.session_state.dados_processo.get('tipo_acao', '')
                ) if st.session_state.dados_processo.get('tipo_acao', '') in ["", "Previdenciária", "Securitária", "Trabalhista", "Cível", "Outras"] else 0
            )
            
            st.session_state.dados_processo['data_nomeacao'] = st.date_input(
                "Data da Nomeação:",
                value=st.session_state.dados_processo.get('data_nomeacao', datetime.now().date())
            )
    
    # Seção: Dados do Autor
    elif opcao == "Dados do Autor":
        st.markdown('<div class="section-header"><h2>👤 Dados do Autor</h2></div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.session_state.dados_autor['nome'] = st.text_input(
                "Nome Completo:",
                value=st.session_state.dados_autor.get('nome', ''),
                help="Nome completo do periciando"
            )
            
            cpf_input = st.text_input(
                "CPF:",
                value=st.session_state.dados_autor.get('cpf', ''),
                help="Digite apenas números"
            )
            
            if cpf_input:
                if validar_cpf(cpf_input):
                    st.session_state.dados_autor['cpf'] = formatar_cpf(cpf_input)
                    st.success("✅ CPF válido")
                else:
                    st.error("❌ CPF inválido")
                    st.session_state.dados_autor['cpf'] = cpf_input
            
            st.session_state.dados_autor['rg'] = st.text_input(
                "RG:",
                value=st.session_state.dados_autor.get('rg', ''),
                help="Número do documento de identidade"
            )
            
            st.session_state.dados_autor['data_nascimento'] = st.date_input(
                "Data de Nascimento:",
                value=st.session_state.dados_autor.get('data_nascimento', datetime.now().date())
            )
        
        with col2:
            st.session_state.dados_autor['profissao'] = st.text_input(
                "Profissão:",
                value=st.session_state.dados_autor.get('profissao', ''),
                help="Profissão do periciando"
            )
            
            st.session_state.dados_autor['endereco'] = st.text_area(
                "Endereço Completo:",
                value=st.session_state.dados_autor.get('endereco', ''),
                help="Endereço completo com CEP"
            )
            
            st.session_state.dados_autor['telefone'] = st.text_input(
                "Telefone:",
                value=st.session_state.dados_autor.get('telefone', ''),
                help="Telefone para contato"
            )
            
            st.session_state.dados_autor['estado_civil'] = st.selectbox(
                "Estado Civil:",
                ["", "Solteiro(a)", "Casado(a)", "Divorciado(a)", "Viúvo(a)", "União Estável"],
                index=["", "Solteiro(a)", "Casado(a)", "Divorciado(a)", "Viúvo(a)", "União Estável"].index(
                    st.session_state.dados_autor.get('estado_civil', '')
                ) if st.session_state.dados_autor.get('estado_civil', '') in ["", "Solteiro(a)", "Casado(a)", "Divorciado(a)", "Viúvo(a)", "União Estável"] else 0
            )
                # Seção: Dados da Perícia
    elif opcao == "Dados da Perícia":
        st.markdown('<div class="section-header"><h2>🔍 Dados da Perícia</h2></div>', unsafe_allow_html=True)
        
        # Dados básicos da perícia
        col1, col2 = st.columns(2)
        
        with col1:
            st.session_state.dados_pericia['data_pericia'] = st.date_input(
                "Data da Perícia:",
                value=st.session_state.dados_pericia.get('data_pericia', datetime.now().date())
            )
            
            st.session_state.dados_pericia['local_pericia'] = st.text_input(
                "Local da Perícia:",
                value=st.session_state.dados_pericia.get('local_pericia', ''),
                help="Endereço onde foi realizada a perícia"
            )
            
            st.session_state.dados_pericia['comparecimento'] = st.selectbox(
                "Comparecimento do Autor:",
                ["Compareceu", "Não compareceu", "Compareceu com atraso"],
                index=["Compareceu", "Não compareceu", "Compareceu com atraso"].index(
                    st.session_state.dados_pericia.get('comparecimento', 'Compareceu')
                )
            )
        
        with col2:
            st.session_state.dados_pericia['acompanhante'] = st.text_input(
                "Acompanhante (se houver):",
                value=st.session_state.dados_pericia.get('acompanhante', ''),
                help="Nome do acompanhante, se presente"
            )
            
            st.session_state.dados_pericia['documentos_apresentados'] = st.text_area(
                "Documentos Apresentados:",
                value=st.session_state.dados_pericia.get('documentos_apresentados', ''),
                help="Liste os documentos médicos apresentados"
            )
        
        # Histórico e Queixas
        st.markdown("### 📋 Histórico e Queixas")
        
        st.session_state.dados_pericia['historico_doenca'] = st.text_area(
            "Histórico da Doença Atual:",
            value=st.session_state.dados_pericia.get('historico_doenca', ''),
            help="Descreva o histórico da condição médica",
            height=150
        )
        
        st.session_state.dados_pericia['queixa_principal'] = st.text_area(
            "Queixa Principal:",
            value=st.session_state.dados_pericia.get('queixa_principal', ''),
            help="Principal queixa relatada pelo periciando",
            height=100
        )
        
        # Exame Físico
        st.markdown("### 🩺 Exame Físico")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.session_state.dados_pericia['estado_geral'] = st.text_area(
                "Estado Geral:",
                value=st.session_state.dados_pericia.get('estado_geral', ''),
                help="Descreva o estado geral do periciando"
            )
            
            st.session_state.dados_pericia['sinais_vitais'] = st.text_area(
                "Sinais Vitais:",
                value=st.session_state.dados_pericia.get('sinais_vitais', ''),
                help="PA, FC, FR, Temperatura, etc."
            )
        
        with col2:
            st.session_state.dados_pericia['exame_fisico_especifico'] = st.text_area(
                "Exame Físico Específico:",
                value=st.session_state.dados_pericia.get('exame_fisico_especifico', ''),
                help="Exame físico direcionado à queixa"
            )
            
            st.session_state.dados_pericia['exames_complementares'] = st.text_area(
                "Análise dos Exames Complementares:",
                value=st.session_state.dados_pericia.get('exames_complementares', ''),
                help="Análise dos exames apresentados"
            )
        
        # Discussão e Conclusão
        st.markdown("### 📝 Discussão e Conclusão")
        
        st.session_state.dados_pericia['discussao'] = st.text_area(
            "Discussão do Caso:",
            value=st.session_state.dados_pericia.get('discussao', ''),
            help="Discussão técnica do caso",
            height=200
        )
        
        st.session_state.dados_pericia['conclusao'] = st.text_area(
            "Conclusão:",
            value=st.session_state.dados_pericia.get('conclusao', ''),
            help="Conclusão pericial",
            height=150
        )
        
        # Capacidade Laborativa
        st.markdown("### 💼 Capacidade Laborativa")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.session_state.dados_pericia['capacidade_laborativa'] = st.selectbox(
                "Capacidade Laborativa:",
                ["", "Preservada", "Parcialmente comprometida", "Totalmente comprometida"],
                index=["", "Preservada", "Parcialmente comprometida", "Totalmente comprometida"].index(
                    st.session_state.dados_pericia.get('capacidade_laborativa', '')
                ) if st.session_state.dados_pericia.get('capacidade_laborativa', '') in ["", "Preservada", "Parcialmente comprometida", "Totalmente comprometida"] else 0
            )
            
            st.session_state.dados_pericia['data_inicio_incapacidade'] = st.date_input(
                "Data de Início da Incapacidade (se houver):",
                value=st.session_state.dados_pericia.get('data_inicio_incapacidade', datetime.now().date())
            )
        
        with col2:
            st.session_state.dados_pericia['grau_incapacidade'] = st.text_input(
                "Grau de Incapacidade (%):",
                value=st.session_state.dados_pericia.get('grau_incapacidade', ''),
                help="Percentual de incapacidade, se aplicável"
            )
            
            st.session_state.dados_pericia['prognóstico'] = st.selectbox(
                "Prognóstico:",
                ["", "Bom", "Regular", "Reservado", "Ruim"],
                index=["", "Bom", "Regular", "Reservado", "Ruim"].index(
                    st.session_state.dados_pericia.get('prognóstico', '')
                ) if st.session_state.dados_pericia.get('prognóstico', '') in ["", "Bom", "Regular", "Reservado", "Ruim"] else 0
            )
                # Seção: Anexos
    elif opcao == "Anexos":
        st.markdown('<div class="section-header"><h2>📎 Anexos</h2></div>', unsafe_allow_html=True)
        
        # Upload da foto do autor
        st.markdown("""
        <div class="upload-section">
            <h3>📷 Foto do Autor (3x4)</h3>
            <p>Faça o upload da foto do periciando que será incluída na identificação do laudo.</p>
        </div>
        """, unsafe_allow_html=True)
        
        foto_autor = st.file_uploader(
            "Selecione a foto do autor:",
            type=['jpg', 'jpeg', 'png'],
            key="foto_autor_upload",
            help="Formatos aceitos: JPG, JPEG, PNG. A foto será redimensionada para 3x4."
        )
        
        if foto_autor is not None:
            st.session_state.foto_autor = foto_autor.read()
            
            # Mostrar preview da foto
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                st.image(st.session_state.foto_autor, caption="Preview da foto do autor", width=200)
                st.success("✅ Foto do autor carregada com sucesso!")
        
        st.markdown("---")
        
        # Upload das fotos dos documentos
        st.markdown("""
        <div class="upload-section">
            <h3>📋 Fotos dos Documentos e Exames</h3>
            <p>Faça o upload das fotos dos documentos médicos, exames e atestados que serão anexados ao laudo.</p>
        </div>
        """, unsafe_allow_html=True)
        
        fotos_documentos = st.file_uploader(
            "Selecione as fotos dos documentos:",
            type=['jpg', 'jpeg', 'png'],
            accept_multiple_files=True,
            key="fotos_documentos_upload",
            help="Formatos aceitos: JPG, JPEG, PNG. Você pode selecionar múltiplas imagens."
        )
        
        if fotos_documentos:
            st.session_state.fotos_documentos = []
            for foto in fotos_documentos:
                st.session_state.fotos_documentos.append(foto.read())
            
            # Mostrar preview das fotos
            st.markdown("### 👀 Preview dos Documentos:")
            
            cols = st.columns(3)
            for i, foto_bytes in enumerate(st.session_state.fotos_documentos):
                with cols[i % 3]:
                    st.image(foto_bytes, caption=f"Documento {i+1}", width=150)
            
            st.success(f"✅ {len(st.session_state.fotos_documentos)} documento(s) carregado(s) com sucesso!")
        
        # Informações sobre os anexos
        if st.session_state.foto_autor or st.session_state.fotos_documentos:
            st.markdown("---")
            st.markdown("### 📊 Resumo dos Anexos:")
            
            info_anexos = []
            if st.session_state.foto_autor:
                info_anexos.append("✅ Foto do autor: Carregada")
            else:
                info_anexos.append("❌ Foto do autor: Não carregada")
            
            if st.session_state.fotos_documentos:
                info_anexos.append(f"✅ Documentos: {len(st.session_state.fotos_documentos)} arquivo(s)")
            else:
                info_anexos.append("❌ Documentos: Nenhum arquivo carregado")
            
            for info in info_anexos:
                st.markdown(f"- {info}")
    
    # Seção: Gerar Laudo
    elif opcao == "Gerar Laudo":
        st.markdown('<div class="section-header"><h2>📄 Gerar Laudo Pericial</h2></div>', unsafe_allow_html=True)
        
        # Verificar se todos os dados necessários estão preenchidos
        dados_completos = True
        campos_obrigatorios = []
        
        # Verificar dados do processo
        if not st.session_state.dados_processo.get('numero_processo'):
            campos_obrigatorios.append("Número do Processo")
            dados_completos = False
        
        # Verificar dados do autor
        if not st.session_state.dados_autor.get('nome'):
            campos_obrigatorios.append("Nome do Autor")
            dados_completos = False
        
        # Verificar dados da perícia
        if not st.session_state.dados_pericia.get('conclusao'):
            campos_obrigatorios.append("Conclusão da Perícia")
            dados_completos = False
        
        if not dados_completos:
            st.warning("⚠️ Alguns campos obrigatórios não foram preenchidos:")
            for campo in campos_obrigatorios:
                st.write(f"- {campo}")
            st.info("💡 Complete os dados nas seções anteriores antes de gerar o laudo.")
        else:
            # Mostrar resumo dos dados
            st.markdown("### 📋 Resumo dos Dados:")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**Processo:**")
                st.write(f"- Número: {st.session_state.dados_processo.get('numero_processo', 'N/A')}")
                st.write(f"- Tipo: {st.session_state.dados_processo.get('tipo_acao', 'N/A')}")
                st.write(f"- Vara: {st.session_state.dados_processo.get('vara', 'N/A')}")
                
                st.markdown("**Autor:**")
                st.write(f"- Nome: {st.session_state.dados_autor.get('nome', 'N/A')}")
                st.write(f"- CPF: {st.session_state.dados_autor.get('cpf', 'N/A')}")
                st.write(f"- Profissão: {st.session_state.dados_autor.get('profissao', 'N/A')}")
            
            with col2:
                st.markdown("**Perícia:**")
                st.write(f"- Data: {st.session_state.dados_pericia.get('data_pericia', 'N/A')}")
                st.write(f"- Comparecimento: {st.session_state.dados_pericia.get('comparecimento', 'N/A')}")
                st.write(f"- Capacidade: {st.session_state.dados_pericia.get('capacidade_laborativa', 'N/A')}")
                
                st.markdown("**Anexos:**")
                st.write(f"- Foto do autor: {'✅' if st.session_state.foto_autor else '❌'}")
                st.write(f"- Documentos: {len(st.session_state.fotos_documentos) if st.session_state.fotos_documentos else 0} arquivo(s)")
            
            st.markdown("---")
            
            # Botão para gerar o laudo
            if st.button("📄 Gerar Laudo Pericial", type="primary", use_container_width=True):
                with st.spinner("Gerando laudo pericial..."):
                    try:
                        doc_bytes = gerar_laudo_completo()
                        
                        if doc_bytes:
                            st.success("✅ Laudo gerado com sucesso!")
                            
                            # Botão de download
                            nome_arquivo = f"Laudo_Pericial_{st.session_state.dados_processo.get('numero_processo', 'SemNumero').replace('/', '_')}.docx"
                            
                            st.download_button(
                                label="📥 Baixar Laudo",
                                data=doc_bytes,
                                file_name=nome_arquivo,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                            
                            st.markdown("""
                            <div class="success-box">
                                <h4>🎉 Laudo Pericial Gerado!</h4>
                                <p>Seu laudo foi gerado com sucesso e está pronto para download. 
                                O documento inclui todas as informações fornecidas e os anexos selecionados.</p>
                            </div>
                            """, unsafe_allow_html=True)
                        else:
                            st.error("❌ Erro ao gerar o laudo. Tente novamente.")
                    
                    except Exception as e:
                        st.error(f"❌ Erro ao gerar o laudo: {str(e)}")
                        def gerar_laudo_completo():
    """
    Gera o documento completo do laudo pericial
    """
    try:
        # Criar documento
        doc = Document()
        
        # Configurar cabeçalho e rodapé
        configurar_cabecalho(doc)
        configurar_rodape(doc)
        
        # Título principal
        adicionar_titulo(doc, "LAUDO PERICIAL MÉDICO", 16, True)
        doc.add_paragraph()
        
        # 1. IDENTIFICAÇÃO DO PROCESSO
        adicionar_secao(doc, "1. IDENTIFICAÇÃO DO PROCESSO", 
                       f"Processo nº: {st.session_state.dados_processo.get('numero_processo', 'N/A')}\n"
                       f"Vara: {st.session_state.dados_processo.get('vara', 'N/A')}\n"
                       f"Comarca: {st.session_state.dados_processo.get('comarca', 'N/A')}\n"
                       f"Juiz(a): {st.session_state.dados_processo.get('juiz', 'N/A')}\n"
                       f"Tipo de Ação: {st.session_state.dados_processo.get('tipo_acao', 'N/A')}\n"
                       f"Data da Nomeação: {st.session_state.dados_processo.get('data_nomeacao', 'N/A')}")
        
        # 2. IDENTIFICAÇÃO DO AUTOR
        texto_identificacao = (f"Nome: {st.session_state.dados_autor.get('nome', 'N/A')}\n"
                              f"CPF: {st.session_state.dados_autor.get('cpf', 'N/A')}\n"
                              f"RG: {st.session_state.dados_autor.get('rg', 'N/A')}\n"
                              f"Data de Nascimento: {st.session_state.dados_autor.get('data_nascimento', 'N/A')}\n"
                              f"Estado Civil: {st.session_state.dados_autor.get('estado_civil', 'N/A')}\n"
                              f"Profissão: {st.session_state.dados_autor.get('profissao', 'N/A')}\n"
                              f"Endereço: {st.session_state.dados_autor.get('endereco', 'N/A')}\n"
                              f"Telefone: {st.session_state.dados_autor.get('telefone', 'N/A')}")
        
        adicionar_secao(doc, "2. IDENTIFICAÇÃO DO AUTOR", texto_identificacao)
        
        # Adicionar foto do autor se disponível
        if st.session_state.foto_autor:
            try:
                img_processada = processar_imagem(st.session_state.foto_autor)
                if img_processada:
                    adicionar_imagem_documento(doc, img_processada, 2.0, 2.7)
            except Exception as e:
                st.warning(f"Não foi possível adicionar a foto do autor: {str(e)}")
        
        # 3. HISTÓRICO
        adicionar_secao(doc, "3. HISTÓRICO", 
                       st.session_state.dados_pericia.get('historico_doenca', 'Não informado'))
        
        # 4. QUEIXA PRINCIPAL
        adicionar_secao(doc, "4. QUEIXA PRINCIPAL", 
                       st.session_state.dados_pericia.get('queixa_principal', 'Não informado'))
        
        # 5. EXAME FÍSICO
        texto_exame = (f"Estado Geral: {st.session_state.dados_pericia.get('estado_geral', 'Não informado')}\n\n"
                      f"Sinais Vitais: {st.session_state.dados_pericia.get('sinais_vitais', 'Não informado')}\n\n"
                      f"Exame Físico Específico: {st.session_state.dados_pericia.get('exame_fisico_especifico', 'Não informado')}")
        
        adicionar_secao(doc, "5. EXAME FÍSICO", texto_exame)
        
        # 6. EXAMES COMPLEMENTARES
        adicionar_secao(doc, "6. ANÁLISE DOS EXAMES COMPLEMENTARES", 
                       st.session_state.dados_pericia.get('exames_complementares', 'Não foram apresentados exames complementares'))
        
        # 7. DISCUSSÃO
        adicionar_secao(doc, "7. DISCUSSÃO", 
                       st.session_state.dados_pericia.get('discussao', 'Não informado'))
        
        # 8. CONCLUSÃO
        adicionar_secao(doc, "8. CONCLUSÃO", 
                       st.session_state.dados_pericia.get('conclusao', 'Não informado'))
        
        # 9. CAPACIDADE LABORATIVA
        texto_capacidade = (f"Capacidade Laborativa: {st.session_state.dados_pericia.get('capacidade_laborativa', 'Não avaliado')}\n"
                           f"Grau de Incapacidade: {st.session_state.dados_pericia.get('grau_incapacidade', 'Não aplicável')}\n"
                           f"Data de Início da Incapacidade: {st.session_state.dados_pericia.get('data_inicio_incapacidade', 'Não aplicável')}\n"
                           f"Prognóstico: {st.session_state.dados_pericia.get('prognóstico', 'Não informado')}")
        
        adicionar_secao(doc, "9. CAPACIDADE LABORATIVA", texto_capacidade)
        
        # Data e assinatura
        doc.add_paragraph()
        data_atual = datetime.now().strftime("%d/%m/%Y")
        adicionar_paragrafo(doc, f"Local e Data: {st.session_state.dados_pericia.get('local_pericia', 'N/A')}, {data_atual}")
        
        doc.add_paragraph()
        doc.add_paragraph()
        adicionar_paragrafo(doc, "_" * 50, WD_ALIGN_PARAGRAPH.CENTER)
        adicionar_paragrafo(doc, "Dr. Hyttallo - Médico Perito", WD_ALIGN_PARAGRAPH.CENTER)
        adicionar_paragrafo(doc, "CRM: [Número do CRM]", WD_ALIGN_PARAGRAPH.CENTER)
        
        # ANEXOS - Nova página
        if st.session_state.fotos_documentos:
            adicionar_quebra_pagina(doc)
            adicionar_titulo(doc, "ANEXOS", 14, True)
            doc.add_paragraph()
            
            for i, foto_bytes in enumerate(st.session_state.fotos_documentos):
                try:
                    img_processada = processar_imagem(foto_bytes, 6, 8)  # Tamanho maior para documentos
                    if img_processada:
                        adicionar_paragrafo(doc, f"Documento {i+1}:", WD_ALIGN_PARAGRAPH.LEFT, 12)
                        adicionar_imagem_documento(doc, img_processada, 6.0, 8.0)
                        doc.add_paragraph()
                except Exception as e:
                    st.warning(f"Não foi possível adicionar o documento {i+1}: {str(e)}")
        
        # Salvar documento em bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
    
    except Exception as e:
        st.error(f"Erro ao gerar documento: {str(e)}")
        return None

# Executar aplicação
if __name__ == "__main__":
    main()
